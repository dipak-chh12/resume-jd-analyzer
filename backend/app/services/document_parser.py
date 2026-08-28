import os
import io
import logging

logger = logging.getLogger("app.services.document_parser")

# Dynamic import fallback for PDF parsing
has_fitz = False
try:
    import fitz  # PyMuPDF
    has_fitz = True
except ImportError:
    logger.warning("PyMuPDF (fitz) not found. Falling back to 'pypdf'.")

has_pypdf = False
try:
    import pypdf
    has_pypdf = True
except ImportError:
    pass

import docx

class DocumentParser:
    @staticmethod
    def _clean_extracted_text(text: str) -> str:
        """Normalize whitespace, fix common PDF font kerning artifacts, and clean up line endings."""
        import re
        if not text:
            return ""
        # Replace carriage returns
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Replace continuous multiple blank lines with max 2 newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Fix spaced single-character kerning sequences (e.g. 'D I P A K' -> 'DIPAK')
        text = re.sub(r"\b([A-Z])\s+([A-Z])\s+([A-Z])\b", r"\1\2\3", text)
        text = re.sub(r"\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b", r"\1\2\3\4", text)
        return text.strip()

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """Parse text content from PDF bytes, prioritizing PyMuPDF block extraction then falling back to pypdf."""
        text_parts = []
        if has_fitz:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    # Extract text blocks preserving reading order
                    blocks = page.get_text("blocks")
                    for b in blocks:
                        block_text = b[4].strip()
                        if block_text:
                            text_parts.append(block_text)
                doc.close()
                combined = "\n\n".join(text_parts)
                return DocumentParser._clean_extracted_text(combined)
            except Exception as e:
                logger.error(f"PyMuPDF block parsing failed: {e}. Trying standard page text / pypdf...")
        
        if has_pypdf:
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_parts.append(extracted)
                combined = "\n\n".join(text_parts)
                return DocumentParser._clean_extracted_text(combined)
            except Exception as e:
                logger.error(f"pypdf parsing failed: {e}")
                raise ValueError(f"Failed to parse PDF document with all parsers: {str(e)}")
        
        raise ImportError("No PDF parser library is installed. Please install 'pymupdf' or 'pypdf'.")

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """Parse text content from DOCX bytes using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            # Parse tables in DOCX
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += "\n" + " | ".join(row_text)
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        """Parse text from plain text file bytes."""
        try:
            try:
                return file_bytes.decode("utf-8").strip()
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1").strip()
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise ValueError(f"Failed to parse plain text file: {str(e)}")

    @classmethod
    def parse_file(cls, filename: str, file_bytes: bytes) -> str:
        """Parse document file content based on file extension."""
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            text = cls.parse_pdf(file_bytes)
        elif ext in [".docx", ".doc"]:
            text = cls.parse_docx(file_bytes)
        elif ext == ".txt":
            text = cls.parse_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")
        
        if not text.strip():
            raise ValueError("The uploaded document contains no readable text.")
            
        return text
